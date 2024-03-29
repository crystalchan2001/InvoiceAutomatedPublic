from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as WAP
from docx.shared import Pt
from datetime import date, datetime
import os

myAddress = "Address\nSeparated\nBy\nWhitespaces"
lineTo = "Who is this for?"
myBankDetails = "Bank Details: \n\nBank Account: \n\nSort Code: \n\nAccount Number: \n\nIBAN: "

today = date.today()
day = today.day
monthNum = today.strftime('%m')
month = today.strftime('%B')
yearShort = today.strftime('%y')
year = today.strftime('%Y')
dateFormat = "%d/%m/%y %H:%M"

class Invoice():                                                                                                                                                                                                                                                                                    

    def __init__(self, lessonsList):
        self.invoiceDoc = Document()
        self.setStyle()
        
        self.lessonsList = lessonsList

        self.writeAddress()
        self.writeHeader()
        self.writeHours()
        self.writeTotal()
        self.writeBankDetails()

        self.saveInvoice()
        self.openInvoice()

    def getPath(self):
        name = f"Monthly Invoice {month} {year}.docx"
        path = f"{name}"
        return path

    def setStyle(self):
        style = self.invoiceDoc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

    def write(self, text):
        para = self.invoiceDoc.add_paragraph(text)
        para.style = self.invoiceDoc.styles['Normal']
        return para

    def getDate(self):
        date = f"{day}/{monthNum}/{yearShort}"
        return "Date: " + date

    def getCode(self):
        code = f"NAME_{monthNum}_{yearShort}"
        return "Invoice: " + code

    def getSubjects(self):
        return "For: English, Maths, Reasoning, Science"
    
    def addLessonEnd(self, lesson):
        lessonDatetime = datetime.strptime(lesson.strip(), dateFormat)
        hour = str(int(lessonDatetime.strftime('%H')) + 1)
        minute = lessonDatetime.strftime('%M')
        lessonEnd = f"-{hour}:{minute}"
        return lesson + lessonEnd

    def getHours(self):
        # lessons will be in format "DD/MM/YY HH:MM"
        lessons = "\n"
        for lesson in self.lessonsList:
            lesson = self.addLessonEnd(lesson)
            lessons += f"\t-{lesson}\n"
        return lessons

    def getTotal(self):
        hours = len(self.lessonsList)
        total = str(hours * 40)
        return "Total: £" + total

    def writeAddress(self):
        addressPar = self.write(myAddress)
        addressPar.alignment = WAP.CENTER

    def writeHeader(self):
        self.write(lineTo)
        self.write(self.getDate())
        self.write(self.getCode())
        self.write(self.getSubjects())
    
    def writeHours(self):
        hoursPar = self.write(f"Hours: {len(self.lessonsList)}")
        hoursPar.add_run(self.getHours())
    
    def writeTotal(self):
        self.write(self.getTotal())

    def writeBankDetails(self):
        self.write(myBankDetails)
    
    def saveInvoice(self):
        path = self.getPath()
        try:
            self.invoiceDoc.save(path)
        except PermissionError:
            print("Check if you have the file open luv")
            doc = Document(path)
            doc.save(path)

    def openInvoice(self):
        os.startfile(self.getPath())

# test input without GUI
# 06/03/24 16:15,06/03/24 17:45,13/03/24 16:15,13/03/24 17:45,20/03/24 16:15,20/03/24 17:45,27/03/24 16:15,27/03/24 17:45

